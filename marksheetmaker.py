import pandas as pd
import re
import docx
import shutil

class MarkSheetMaker:

    colnum_name = 1
    colnum_number = 0
    colnum_group = None
    num_paragraph_student_num = 1
    num_paragraph_student_module = 3
    
    def __init__(self
            ,module_code
            ,student_list_filename
            ,marksheet_filename
            ,student_list_grouping_var=None):

        self.module_code = module_code
        self.grouping_var = student_list_grouping_var

        # if there is a grouping variable in the first column
        # shift the columns
        if self.grouping_var != None:
            self.colnum_group = 0
            self.colnum_name = self.colnum_name + 1
            self.colnum_number = self.colnum_number + 1

        self.marksheet_filename = marksheet_filename
        self.read_student_list(student_list_filename)
        self.copy_marksheets()

        
    # read in the list of students and modify 
    # columns so that we end up with student number, name, surname
    def read_student_list(self,filename):
        self.student_df = pd.read_csv(
                filepath_or_buffer=filename
                ,header=None)

        self.keep_name_number_cols()
        self.edit_student_number_column()
        self.edit_student_name_column()

    def keep_name_number_cols(self):

        col_list = [self.colnum_number,self.colnum_name]
        col_names = ["number","name"]

        # if there is a grouping var add it to the list
        if type(self.colnum_group) == type(2):
            col_list = [self.colnum_group] + col_list
            col_names  = ["group"] + col_names

        self.student_df = self.student_df[col_list]
        self.student_df.columns = col_names

    # remove trailing slash in student num
    def transform_number(self,x):
        return(re.sub(pattern=r"\/.*",repl="", string=x))

    def edit_student_number_column(self):
        self.student_df["number"] = self.student_df["number"].apply(self.transform_number)

    def capitalize_surname(self,surname):
        surname_split = re.split(pattern=r"\s+",string=surname)

        surname = ""

        surname_split_len = len(surname_split)

        for i, sub_surname in enumerate(surname_split):
            surname += sub_surname.strip().capitalize() 

            if i < surname_split_len - 1:
                surname += " "

        return(surname)

    # remove capitalized surname
    def transform_name(self,x):
        name_split = re.split(pattern=r",\s+", string=x)
        surname = self.capitalize_surname(name_split[0])
        name = name_split[1]
        
        return(name + " " + surname)

    def edit_student_name_column(self):
        self.student_df["name"] = self.student_df["name"].apply(self.transform_name)

    def make_marksheet_path(self, row, path_end):

        path = ""

        # if there is a grouping variable add the dir
        if self.grouping_var != None:
            path = self.group_folder_prefix + getattr(row,"group") + "/"

        path += getattr(row, "number") + \
            " " + getattr(row, "name") + \
            " " + path_end

        return(path)

    def edit_marksheet(self, path, row):
        
        marksheet_doc = docx.Document(docx=path)

        marksheet_doc.paragraphs[self.num_paragraph_student_num].text = "Student number: " + getattr(row,"number") + " "*10 + "MARK:"
        
        marksheet_doc.paragraphs[self.num_paragraph_student_module].text = "Module: " + self.module_code

        marksheet_doc.save(path_or_stream=path)

    def check_group_folder_exists(self, row):
        # make folder name
        folder_name = self.group_folder_prefix + getattr(row, "group")
        if !os.path.exists(folder_name):
            os.mkdir(folder_name)

    def copy_marksheets(self):
        for row in self.student_df.itertuples():

            if self.grouping_var != None:
                check_group_folder_exists(row)

            new_path_marksheet = self.make_marksheet_path(row=row, path_end=self.marksheet_filename)
            shutil.copy(self.marksheet_filename, new_path_marksheet)
            
            self.edit_marksheet(path=new_path_marksheet, row=row)

msm = MarkSheetMaker(module_code = "BIO3148"
        ,student_list_filename="student_list.csv"
        ,marksheet_filename="CLES Penryn Oral Presentation Feedback Sheet and marking criteria AMcG.docx")
